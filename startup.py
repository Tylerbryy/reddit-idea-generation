import praw
import openai
import json
import os
from dotenv import load_dotenv
from halo import Halo
from docx import Document
from openai import OpenAI

# Load API keys from .env file
load_dotenv()

# Reddit API setup
reddit = praw.Reddit(
    client_id=os.getenv("REDDIT_CLIENT_ID"),
    client_secret=os.getenv("REDDIT_CLIENT_SECRET"),
    user_agent="script by @tylergibbs"
)

# OpenAI API setup
openai.api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
# Function to fetch posts from a subreddit
def fetch_posts(subreddit_name, post_type="hot", limit=15):
    subreddit = reddit.subreddit(subreddit_name)
    fetched_posts = []
    processed_post_ids = load_processed_posts(f"processed_posts/{subreddit_name}_{post_type}.json")

    # Generator expressions for different post types
    post_generators = {
        "hot": subreddit.hot(limit=None),
        "new": subreddit.new(limit=None),
        "top": subreddit.top(limit=None)
    }

    try:
        posts = post_generators[post_type]
    except KeyError:
        raise ValueError("Invalid post type specified. Choose 'hot', 'new', or 'top'.")

    for post in posts:
        if post.id not in processed_post_ids and len(fetched_posts) < limit:
            fetched_posts.append(post)
        if len(fetched_posts) == limit:
            break

    return fetched_posts

# Function to load processed posts from a JSON file
def load_processed_posts(filename):
    try:
        with open(f"processed_posts/{filename}", 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        return []

# Function to save processed posts to a JSON file
def save_processed_posts(processed_posts, filename):
    os.makedirs("processed_posts", exist_ok=True)
    with open(f"processed_posts/{filename}", 'w', encoding='utf-8') as f:
        json.dump(processed_posts, f, ensure_ascii=False, indent=4)

# Function to generate startup ideas using OpenAI with retries
@Halo(text='Processing', spinner='dots')
def generate_idea(post_title, post_content):
    retries = 0
    while retries < 10:  # Maximum 10 retries
        try:
            messages = []
            messages.append({"role": "user", "content": f"Here is a Reddit post titled '{post_title}' with the content: '{post_content}'. Brainstorm some video ideas of what would be a thought-provoking video title and outline for my YouTube channel?"})
            
            response = client.chat.completions.create(
                model="gpt-4-1106-preview",
                messages=messages
            )
            
            return response.choices[0].message.content
        except Exception as e:
            print(f"API Error: {e}")
            retries += 1
            if retries >= 10:
                return None
            
# Main script execution
subreddit_name = input("Enter the subreddit name: ")
post_type = input("Enter the type of posts to fetch (hot, new, top): ").lower()

def save_to_docx(ideas, filename=f"ideas_{subreddit_name}_{post_type}.docx"):
    doc = Document()
    doc.add_heading('Ideas', 0)

    for idea in ideas:
        doc.add_heading(idea['post_title'], level=1)
        doc.add_paragraph(idea['idea'])

    doc.save(filename)
    print(f"Startup ideas have been saved to '{filename}'.")

# Load processed posts
processed_posts_filename = f"processed_posts_{subreddit_name}_{post_type}.json"
processed_posts = load_processed_posts(processed_posts_filename)

try:
    posts = fetch_posts(subreddit_name, post_type=post_type)
except ValueError as e:
    print(e)
    exit()

ideas = []
for post in posts:
    if post.id in processed_posts:
        print(f"Skipping already processed post: {post.title}")
        continue  # Skip this post as it has already been processed

    print(f"Analyzing post: {post.title}")
    idea = generate_idea(post.title, post.selftext)
    if idea:
        ideas.append({
            "post_title": post.title,
            "idea": idea
        })
        processed_posts.append(post.id)  # Add the post ID to the processed list
    else:
        print("No viable startup idea found or API limit reached for this post.\n")

# Save the updated list of processed posts
save_processed_posts(processed_posts, processed_posts_filename)

# Ask user for preferred output format
output_format = input("Choose the output format (docx or json): ").lower()

if output_format == "docx":
    save_to_docx(ideas)
elif output_format == "json":
    # Output results to a JSON file
    os.makedirs("ideas", exist_ok=True)
    with open(f'ideas/ideas_{subreddit_name}_{post_type}.json', 'w', encoding='utf-8') as f:
        json.dump(ideas, f, ensure_ascii=False, indent=4)
    print(f"Startup ideas have been saved to 'ideas/ideas_{subreddit_name}_{post_type}.json'.")
else:
    print("Invalid output format selected. Please choose either 'docx' or 'json'.")
if output_format == "docx":
    save_to_docx(ideas)
elif output_format == "json":
    # Output results to a JSON file
    os.makedirs("ideas", exist_ok=True)
    with open(f'ideas/ideas_{subreddit_name}_{post_type}.json', 'w', encoding='utf-8') as f:
        json.dump(ideas, f, ensure_ascii=False, indent=4)
    print("Startup ideas have been saved to 'ideas.json'.")
else:
    print("Invalid output format selected. Please choose either 'docx' or 'json'.")
